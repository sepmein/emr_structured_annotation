from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


OUTPUT = Path(
    "/Users/sepmein/Documents/EMR labelling/outputs/"
    "关于2027年科普宣传类经费情况的汇报.docx"
)

TITLE_FONT = "FZXiaoBiaoSong-B05S"
BODY_FONT = "FangSong_GB2312"
H1_FONT = "SimHei"
H2_FONT = "KaiTi_GB2312"


def set_run_font(run, name, size_pt, *, bold=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)


def set_paragraph_geometry(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line=True,
    keep_with_next=False,
):
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(28)
    fmt.first_line_indent = Pt(32) if first_line else Pt(0)
    fmt.keep_with_next = keep_with_next
    fmt.widow_control = True


def add_body_paragraph(doc, text):
    paragraph = doc.add_paragraph(style="正文")
    set_paragraph_geometry(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, BODY_FONT, 16)
    return paragraph


def add_h1(doc, text):
    paragraph = doc.add_paragraph(style="一级标题")
    set_paragraph_geometry(paragraph, keep_with_next=True)
    run = paragraph.add_run(text)
    set_run_font(run, H1_FONT, 16)
    return paragraph


def add_h2_body(doc, heading, body):
    paragraph = doc.add_paragraph(style="正文")
    set_paragraph_geometry(paragraph)
    heading_run = paragraph.add_run(heading)
    set_run_font(heading_run, H2_FONT, 16, bold=True)
    body_run = paragraph.add_run(body)
    set_run_font(body_run, BODY_FONT, 16)
    return paragraph


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    left = paragraph.add_run("— ")
    set_run_font(left, "Songti SC", 14)

    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), "Songti SC")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "28")
    rpr.append(rfonts)
    rpr.append(size)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(rpr)
    run.append(text)
    field.append(run)
    paragraph._p.append(field)

    right = paragraph.add_run(" —")
    set_run_font(right, "Songti SC", 14)


doc = Document()
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.orientation = 0
section.top_margin = Mm(37)
section.bottom_margin = Mm(35)
section.left_margin = Mm(28)
section.right_margin = Mm(26)
section.header_distance = Mm(15)
section.footer_distance = Mm(17)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(16)
normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
normal.paragraph_format.line_spacing = Pt(28)

for style_name, font_name in (
    ("正文", BODY_FONT),
    ("一级标题", H1_FONT),
):
    if style_name not in styles:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles[style_name]
    style.font.name = font_name
    style.font.size = Pt(16)
    style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(28)

footer = section.footer
footer.is_linked_to_previous = False
footer_paragraph = footer.paragraphs[0]
add_page_number(footer_paragraph)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(28)
title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
title.paragraph_format.line_spacing = Pt(32)
title.paragraph_format.keep_with_next = True
title_run = title.add_run("关于2027年科普宣传类经费情况的汇报")
set_run_font(title_run, TITLE_FONT, 22)

add_body_paragraph(
    doc,
    "根据财务处有关通知，现将监测预警所2027年科普宣传类经费预算、"
    "实施依据、产出绩效、实施方案及较上年变化情况汇报如下。",
)

add_h1(doc, "一、实施依据")
add_body_paragraph(
    doc,
    "根据《中华人民共和国传染病防治法》《传染病疫情风险评估管理办法"
    "（试行）》等文件，并结合《传染病疫情预警管理办法（试行）》有关要求，"
    "疾病预防控制机构应当根据传染病疫情风险评估结果，及时向社会发布健康"
    "风险提示并开展健康宣教。",
)
add_body_paragraph(
    doc,
    "《传染病疫情预警管理办法（试行）》第四条明确，健康风险提示是指由"
    "各级疾病预防控制机构向公众提示传染病疫情风险，帮助公众提高对传染病"
    "的科学认识，引导采取必要、适当的防护措施，保护自身健康。结合我所传染病"
    "疫情风险评估、健康风险提示发布和公众健康教育工作需要，通过制作科普用品"
    "和科普视频，进一步丰富健康风险提示传播载体，提高风险沟通的科学性、"
    "针对性和覆盖面。",
)

add_h1(doc, "二、经费安排及测算依据")
add_body_paragraph(doc, "2027年科普宣传类经费预算共计6万元，具体包括以下两项：")
add_h2_body(
    doc,
    "（一）传染病预防控制健康风险提示科普用品制作经费4万元。",
    "结合风险评估和健康风险提示工作需要，制作健康风险提示科普用品，用于"
    "重点传染病防控宣传活动和公众健康教育。按照全市16个区、每区50份测算，"
    "共制作800份，单价50元，预算测算依据明确。",
)
add_h2_body(
    doc,
    "（二）科普视频策划制作经费2万元。",
    "在全球疫情常态化背景下，为提高入境人员健康科普的精准性和覆盖面，计划"
    "购买第三方专业服务，制作1个约3分钟的健康科普视频。视频将围绕入境人员"
    "重点传染病防护知识和健康提示开展内容策划与制作，兼顾科学性、通俗性和"
    "国际传播需求。预算根据市场标准和实际制作需求综合测算，金额合理合规。",
)

add_h1(doc, "三、实施方案")
add_h2_body(
    doc,
    "（一）科普用品制作。",
    "以传染病疫情风险评估结果和健康风险提示内容为基础，根据年度重点传染病"
    "防控任务确定宣传主题，依次开展内容设计、专业审核、采购制作、成品验收和"
    "分发使用。制作完成后，结合健康风险提示发布、重点宣传活动和公众健康教育，"
    "向全市16个区进行配置和发放。",
)
add_h2_body(
    doc,
    "（二）科普视频制作。",
    "根据入境人员健康科普需求，按照主题确定、资料整理、脚本策划、专业审核、"
    "视觉设计与拍摄制作、成片审看、项目验收和宣传使用等环节推进。通过规范"
    "采购程序委托具备专业策划和视频制作能力的第三方机构实施，明确各阶段任务"
    "和质量要求，确保按期产出内容科学准确、表达清晰易懂、具有较好传播效果的"
    "健康科普作品。",
)

add_h1(doc, "四、产出绩效")
add_body_paragraph(
    doc,
    "项目实施后，预计完成健康风险提示科普用品制作800份、健康科普视频1个。"
    "科普用品将服务全市健康风险提示工作，用于重点宣传活动和公众健康教育；"
    "科普视频将面向入境相关人群传播重点传染病防护和健康提示信息，提升健康科普"
    "的精准性、覆盖面和传播效果。通过线上线下科普载体协同使用，促进专业风险"
    "信息向公众可理解、可接受、可采取行动的健康信息转化，进一步提升传染病"
    "防控知识知晓率和风险防范意识。",
)

add_h1(doc, "五、较上年主要变化")
add_body_paragraph(
    doc,
    "2027年新增传染病预防控制健康风险提示科普用品制作预算4万元，主要用于"
    "配合健康风险提示常态化发布、重点传染病防控宣传和公众健康教育工作。科普"
    "视频制作经费仍为2万元，与2026年相比金额和产出数量保持不变。",
)
add_body_paragraph(
    doc,
    "科普宣传类经费合计由2026年的2万元增加至2027年的6万元，增加4万元。"
    "宣传方式由以视频传播为主，进一步拓展为“数字视频传播与线下科普用品覆盖"
    "相结合”，增强健康风险提示在全市各区的传播覆盖和实际应用效果。",
)

add_body_paragraph(doc, "以上汇报。")

signature = doc.add_paragraph()
signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
signature.paragraph_format.space_before = Pt(28)
signature.paragraph_format.space_after = Pt(0)
signature.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
signature.paragraph_format.line_spacing = Pt(28)
signature_run = signature.add_run("监测预警所")
set_run_font(signature_run, BODY_FONT, 16)

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
date.paragraph_format.space_before = Pt(0)
date.paragraph_format.space_after = Pt(0)
date.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
date.paragraph_format.line_spacing = Pt(28)
date_run = date.add_run("2026年7月30日")
set_run_font(date_run, BODY_FONT, 16)

doc.core_properties.title = "关于2027年科普宣传类经费情况的汇报"
doc.core_properties.subject = "监测预警所科普宣传类经费情况汇报"
doc.core_properties.author = "监测预警所"

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
